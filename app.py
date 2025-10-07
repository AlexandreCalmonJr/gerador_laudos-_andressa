from flask import Flask, render_template, request, send_from_directory, flash, redirect, url_for, jsonify
from docx import Document
from docx.shared import Inches
import os
import uuid
import logging
import time
from werkzeug.utils import secure_filename
from datetime import datetime

# Configuração de logging aprimorada
logging.basicConfig(
    filename='app_errors.log', 
    level=logging.INFO,
    format='%(asctime)s [%(levelname)s] %(name)s: %(message)s'
)

app = Flask(__name__)
app.secret_key = os.environ.get('SECRET_KEY', 'uma-chave-secreta-muito-forte-e-dificil')

# Configuração das pastas
UPLOAD_FOLDER = 'uploads'
GENERATED_FOLDER = 'gerados'
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif', 'bmp', 'webp'}
MAX_FILE_SIZE = 16 * 1024 * 1024  # 16MB

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(GENERATED_FOLDER, exist_ok=True)

app.config.update(
    UPLOAD_FOLDER=UPLOAD_FOLDER,
    GENERATED_FOLDER=GENERATED_FOLDER,
    MAX_FILE_AGE_SECONDS=24 * 3600,
    MAX_CONTENT_LENGTH=MAX_FILE_SIZE
)

def allowed_file(filename):
    """Verifica se a extensão do arquivo é permitida"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def validar_imagem(arquivo):
    """Valida o arquivo de imagem"""
    if not arquivo or arquivo.filename == '':
        return False, "Nenhum arquivo selecionado"
    
    if not allowed_file(arquivo.filename):
        return False, f"Tipo de arquivo não permitido. Use: {', '.join(ALLOWED_EXTENSIONS)}"
    
    return True, None

def gerar_documento(dados, nome_arquivo_saida):
    """Gera o documento Word com os dados fornecidos"""
    try:
        template_path = 'Vistoria_Modelo.docx'
        if not os.path.exists(template_path):
            app.logger.error(f"Template não encontrado: {template_path}")
            flash("Erro: Template do documento não encontrado no servidor.", "error")
            return False

        doc = Document(template_path)
        
        # Substituição de texto em todos os parágrafos
        all_paragraphs = list(doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_paragraphs.extend(cell.paragraphs)

        for p in all_paragraphs:
            for key, value in dados['texto'].items():
                if key in p.text:
                    for run in p.runs:
                        if key in run.text:
                            run.text = run.text.replace(key, str(value))
        
        # Inserção de imagens
        imagens_processadas = 0
        imagens_falhas = 0
        
        for p in doc.paragraphs:
            for comodo_id, imagens in dados['imagens'].items():
                marcador_imagem = f"{{{{IMAGENS_{comodo_id.upper()}}}}}"
                if marcador_imagem in p.text:
                    p.text = ''
                    for img_path in imagens:
                        try:
                            if os.path.exists(img_path):
                                p.add_run().add_picture(img_path, width=Inches(3.0))
                                imagens_processadas += 1
                            else:
                                app.logger.warning(f"Imagem não encontrada: {img_path}")
                                imagens_falhas += 1
                        except Exception as e:
                            app.logger.error(f"Erro ao adicionar imagem {img_path}: {e}")
                            imagens_falhas += 1

        caminho_completo = os.path.join(app.config['GENERATED_FOLDER'], nome_arquivo_saida)
        doc.save(caminho_completo)
        
        app.logger.info(f"Documento gerado: {nome_arquivo_saida} ({imagens_processadas} imagens, {imagens_falhas} falhas)")
        
        if imagens_falhas > 0:
            flash(f"Documento gerado com sucesso! ({imagens_falhas} imagens não puderam ser processadas)", "warning")
        else:
            flash("Documento gerado com sucesso!", "success")
        
        return True
        
    except Exception as e:
        app.logger.error(f"Erro ao gerar documento: {str(e)}", exc_info=True)
        flash("Erro ao gerar o documento. Tente novamente.", "error")
        return False

def limpar_arquivos_antigos():
    """Remove arquivos antigos das pastas de upload e gerados"""
    now = time.time()
    max_age = app.config['MAX_FILE_AGE_SECONDS']
    arquivos_removidos = 0
    
    for folder in [app.config['UPLOAD_FOLDER'], app.config['GENERATED_FOLDER']]:
        if not os.path.exists(folder):
            continue
            
        for filename in os.listdir(folder):
            file_path = os.path.join(folder, filename)
            if os.path.isfile(file_path):
                try:
                    file_age = now - os.path.getmtime(file_path)
                    if file_age > max_age:
                        os.remove(file_path)
                        arquivos_removidos += 1
                except Exception as e:
                    app.logger.error(f"Erro ao remover arquivo {file_path}: {e}")
    
    if arquivos_removidos > 0:
        app.logger.info(f"Limpeza automática: {arquivos_removidos} arquivo(s) removido(s)")

def limpar_uploads_sessao(caminhos):
    """Remove os arquivos de upload da sessão atual"""
    for path in caminhos:
        try:
            if os.path.exists(path):
                os.remove(path)
        except Exception as e:
            app.logger.error(f"Erro ao remover upload temporário {path}: {e}")

# ROTAS
@app.route('/')
def index():
    limpar_arquivos_antigos()
    return render_template('form.html')

@app.route('/gerar', methods=['POST'])
def gerar_laudo():
    try:
        # Validação dos campos obrigatórios
        campos_obrigatorios = ['LOCATARIO_NOME_1', 'ENDERECO_IMOVEL', 'DATA_VISTORIA']
        for campo in campos_obrigatorios:
            if campo not in request.form or not request.form[campo].strip():
                flash(f"O campo '{campo.replace('_', ' ').title()}' é obrigatório.", "error")
                return redirect(url_for('index'))

        # Coleta de dados do formulário
        dados_texto = {f"{{{{{key}}}}}": value.strip() for key, value in request.form.items() if value.strip()}
        dados_imagens = {}
        caminhos_temporarios = []
        
        comodos = ["externa", "sala", "cozinha", "banheiro_social", "quarto1", "quarto2", 
                   "servico", "fundos", "gourmet", "banheiro_edicula", "laterais"]
        
        # Processamento de imagens
        for comodo in comodos:
            arquivos = request.files.getlist(f'imagens_{comodo}')
            caminhos_salvos = []
            
            for arquivo in arquivos:
                valido, mensagem_erro = validar_imagem(arquivo)
                
                if not valido:
                    if mensagem_erro != "Nenhum arquivo selecionado":
                        app.logger.warning(f"Imagem inválida em {comodo}: {mensagem_erro}")
                    continue
                
                # Salva o arquivo com nome seguro
                filename = secure_filename(arquivo.filename)
                nome_unico = f"{uuid.uuid4().hex[:8]}_{filename}"
                caminho = os.path.join(app.config['UPLOAD_FOLDER'], nome_unico)
                
                arquivo.save(caminho)
                caminhos_salvos.append(caminho)
                caminhos_temporarios.append(caminho)
            
            if caminhos_salvos:
                dados_imagens[comodo] = caminhos_salvos

        dados_completos = {'texto': dados_texto, 'imagens': dados_imagens}

        # Gera nome único para o arquivo
        nome_locatario = secure_filename(request.form.get('LOCATARIO_NOME_1').replace(' ', '_'))
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        nome_arquivo_final = f"Laudo_Vistoria_{nome_locatario}_{timestamp}.docx"
        
        # Gera o documento
        sucesso = gerar_documento(dados_completos, nome_arquivo_final)

        # Limpa uploads temporários
        limpar_uploads_sessao(caminhos_temporarios)
        
        if sucesso:
            return redirect(url_for('resultado', filename=nome_arquivo_final))
        else:
            return redirect(url_for('index'))
            
    except Exception as e:
        app.logger.error(f"Erro no processamento do laudo: {str(e)}", exc_info=True)
        flash("Erro inesperado ao processar o formulário. Tente novamente.", "error")
        return redirect(url_for('index'))

@app.route('/resultado/<filename>')
def resultado(filename):
    # Verifica se o arquivo existe
    filepath = os.path.join(app.config['GENERATED_FOLDER'], filename)
    if not os.path.exists(filepath):
        flash("Arquivo não encontrado.", "error")
        return redirect(url_for('index'))
    
    return render_template('resultado.html', filename=filename)

@app.route('/download/<filename>')
def download_file(filename):
    try:
        return send_from_directory(
            app.config["GENERATED_FOLDER"], 
            filename, 
            as_attachment=True,
            download_name=filename
        )
    except Exception as e:
        app.logger.error(f"Erro ao fazer download do arquivo {filename}: {e}")
        flash("Erro ao fazer download do arquivo.", "error")
        return redirect(url_for('index'))

@app.route('/health')
def health_check():
    """Endpoint para verificação de saúde da aplicação"""
    return jsonify({
        'status': 'healthy',
        'timestamp': datetime.now().isoformat()
    }), 200

@app.errorhandler(413)
def request_entity_too_large(error):
    flash(f"Arquivo muito grande! O tamanho máximo permitido é {MAX_FILE_SIZE // (1024*1024)}MB.", "error")
    return redirect(url_for('index'))

@app.errorhandler(500)
def internal_error(error):
    app.logger.error(f"Erro interno do servidor: {str(error)}", exc_info=True)
    flash("Erro interno do servidor. Por favor, tente novamente mais tarde.", "error")
    return redirect(url_for('index'))

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)